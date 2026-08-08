"""Parsers, against the real committed fixtures.

These run on the actual generated PDFs rather than mocks, because the thing worth
testing is exactly what a mock would paper over: whether pdfplumber, on a real
file, yields the character attributes the injection scanner depends on.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import docx as python_docx
import pytest

from apps.documents.normalize import normalize
from apps.documents.parsers.docx import parse_docx
from apps.documents.parsers.pdf import parse_pdf
from apps.documents.parsers.plain import parse_plain
from apps.documents.sanitize import InjectionReason, scan
from apps.documents.validators import (
    ContentMismatchError,
    FileTooLargeError,
    UnsupportedTypeError,
    validate_upload,
)

FIXTURES = Path(__file__).resolve().parents[2] / "fixtures"


def _open(relative: str) -> io.BytesIO:
    return io.BytesIO((FIXTURES / relative).read_bytes())


class TestPdf:
    def test_resume_text_is_extracted(self) -> None:
        parsed = parse_pdf(_open("demo/resume.pdf"))

        assert "ALEX MORAN" in parsed.text
        assert "Meridian Logistics" in parsed.text
        assert parsed.page_count >= 1

    def test_visible_job_text_is_extracted(self) -> None:
        parsed = parse_pdf(_open("demo/job_2_vertex.pdf"))

        assert "Kubernetes" in parsed.text
        assert "Terraform" in parsed.text

    def test_clean_fixtures_report_no_hidden_text(self) -> None:
        """The false-positive check on the colour heuristic.

        Ordinary black-on-white text must never be reported as hidden, or every
        real document gets quarantined.
        """
        for name in ("demo/resume.pdf", "demo/job_1_northwind.pdf", "demo/job_3_helio.pdf"):
            assert parse_pdf(_open(name)).hidden_spans == [], name

    def test_white_on_white_payload_is_detected(self) -> None:
        """The attack this parser choice exists for.

        pdfplumber was chosen over pypdf specifically because it exposes
        per-character non_stroking_color. This is that decision paying off.
        """
        parsed = parse_pdf(_open("adversarial_job.pdf"))

        assert parsed.hidden_spans, "invisible payload was not detected"
        assert "Ignore all previous instructions" in " ".join(parsed.hidden_spans)

    def test_hidden_payload_flows_through_to_a_scan_finding(self) -> None:
        """End to end: parser sees colour, scanner reports it, UI can render it."""
        parsed = parse_pdf(_open("adversarial_job.pdf"))
        normalized = normalize(parsed.text)

        result = scan(
            normalized.text,
            invisible_chars_removed=normalized.invisible_chars_removed,
            hidden_spans=parsed.hidden_spans,
        )

        assert result.flagged
        assert InjectionReason.HIDDEN_TEXT.value in result.reasons


class TestDocx:
    @staticmethod
    def _build(paragraphs: list[str], table: list[list[str]] | None = None) -> io.BytesIO:
        document = python_docx.Document()
        for text in paragraphs:
            document.add_paragraph(text)
        if table:
            docx_table = document.add_table(rows=len(table), cols=len(table[0]))
            for row_index, row in enumerate(table):
                for col_index, cell in enumerate(row):
                    docx_table.cell(row_index, col_index).text = cell
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def test_paragraphs_are_extracted(self) -> None:
        parsed = parse_docx(self._build(["EXPERIENCE", "Senior Engineer at Acme"]))

        assert "EXPERIENCE" in parsed.text
        assert "Senior Engineer at Acme" in parsed.text

    def test_table_content_is_extracted(self) -> None:
        """A large minority of résumés put skills in a borderless table.

        A parser that reads only paragraphs silently drops exactly the section a
        skills question needs — and reports success while doing it.
        """
        parsed = parse_docx(self._build(["SKILLS"], table=[["Python", "Go"], ["SQL", "Kafka"]]))

        for skill in ("Python", "Go", "SQL", "Kafka"):
            assert skill in parsed.text, skill

    def test_document_order_is_preserved_across_paragraphs_and_tables(self) -> None:
        """Section detection depends on a heading being adjacent to its content.

        python-docx exposes .paragraphs and .tables as separate flat lists; using
        those would detach the skills table from the SKILLS heading above it.
        """
        parsed = parse_docx(self._build(["SKILLS"], table=[["Python"]]))

        assert parsed.text.index("SKILLS") < parsed.text.index("Python")


class TestPlain:
    def test_utf8(self) -> None:
        assert "café" in parse_plain(io.BytesIO("Résumé for café".encode())).text

    def test_cp1252_fallback(self) -> None:
        """Windows-exported .txt is routinely cp1252; rejecting it would be rude.

        Smart quotes are the giveaway: U+201C/U+201D encode to bytes 0x93/0x94,
        which are not valid UTF-8, so a strict utf-8 decode raises on a file that
        is perfectly readable.
        """
        raw = "Smart “quotes”".encode("cp1252")
        assert b"\x93" in raw  # the byte that breaks a strict utf-8 decode

        parsed = parse_plain(io.BytesIO(raw))

        assert "quotes" in parsed.text


class TestValidation:
    def test_oversize_is_rejected_before_parsing(self) -> None:
        with pytest.raises(FileTooLargeError):
            validate_upload(
                filename="a.pdf", size_bytes=11 * 1024 * 1024, handle=io.BytesIO(b"%PDF-")
            )

    def test_unsupported_extension(self) -> None:
        with pytest.raises(UnsupportedTypeError):
            validate_upload(filename="a.exe", size_bytes=10, handle=io.BytesIO(b"MZ"))

    def test_extension_spoofing_is_rejected(self) -> None:
        """Content-Type is attacker-controlled; the magic bytes are not.

        Renaming an executable to .pdf must not hand it to the PDF parser.
        """
        with pytest.raises(ContentMismatchError):
            validate_upload(
                filename="payload.pdf", size_bytes=64, handle=io.BytesIO(b"MZ\x90\x00" * 16)
            )

    def test_a_real_pdf_passes(self) -> None:
        handle = _open("demo/resume.pdf")

        validated = validate_upload(
            filename="resume.pdf", size_bytes=len(handle.getvalue()), handle=handle
        )

        assert validated.extension == "pdf"

    def test_zip_bomb_ratio_is_rejected(self) -> None:
        """A DOCX is a zip. A 200x ratio is not a résumé."""
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", "A" * 5_000_000)
        buffer.seek(0)

        with pytest.raises(Exception, match=r"malformed|read"):
            validate_upload(filename="bomb.docx", size_bytes=len(buffer.getvalue()), handle=buffer)
