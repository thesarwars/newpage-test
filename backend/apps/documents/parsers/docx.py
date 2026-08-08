"""DOCX extraction.

Tables are extracted, not skipped: a large minority of résumés lay out their
skills or dates in a borderless table, and a parser that reads only paragraphs
silently drops exactly the section a skills question needs.
"""

from __future__ import annotations

from typing import BinaryIO

import docx
import structlog
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from apps.documents.parsers.base import ParsedDocument
from apps.documents.validators import ParseFailedError

log = structlog.get_logger(__name__)

# DOCX has no page concept until it is rendered. Estimated from content so the
# page cap means something, rather than reported as a lie.
_CHARS_PER_PAGE_ESTIMATE = 1800


def parse_docx(handle: BinaryIO) -> ParsedDocument:
    handle.seek(0)
    try:
        document = docx.Document(handle)
        text = "\n".join(_iter_block_text(document))
    except Exception as exc:
        log.info("docx_parse_failed", exc_type=type(exc).__name__)
        raise ParseFailedError() from exc

    return ParsedDocument(
        text=text,
        page_count=max(1, len(text) // _CHARS_PER_PAGE_ESTIMATE + 1),
        warnings=["page_count_is_estimated"],
    )


def _iter_block_text(document: DocxDocument) -> list[str]:
    """Walk paragraphs and tables in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate flat lists, which
    loses their interleaving — a skills table would end up detached from the
    "Skills" heading that introduces it, and section detection depends on that
    adjacency. Walking the underlying XML body preserves the real order.
    """
    blocks: list[str] = []
    body = document.element.body

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, document).text.strip()
            if text:
                blocks.append(text)
        elif child.tag == qn("w:tbl"):
            blocks.extend(_table_lines(Table(child, document)))

    return blocks


def _table_lines(table: Table) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        # Deduplicate horizontally merged cells, which python-docx repeats.
        deduped: list[str] = []
        for cell in cells:
            if not deduped or deduped[-1] != cell:
                deduped.append(cell)
        line = "  ".join(cell for cell in deduped if cell)
        if line:
            lines.append(line)
    return lines
